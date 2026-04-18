from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "output" / "doc"
TMP_DIR = ROOT / "tmp" / "docs"
DOCX_PATH = OUTPUT_DIR / "无人机编队中通信技术应用报告.docx"
FIG_PATH = TMP_DIR / "uav_angle_comm_topology.png"


TITLE = "无人机编队中通信技术应用报告"

ABSTRACT = (
    "无人机编队正由单纯的队形保持逐步走向协同侦察、目标围捕、应急通信中继和复杂环境作业，"
    "通信系统已成为决定编队控制性能的重要基础。对于围绕角度信息或角度约束开展的编队控制方案，"
    "通信不仅承担状态交换功能，还直接影响相对方位估计、队形刚性维持、拓扑切换稳定性以及控制律实时实现。"
    "本文围绕无人机编队任务中的通信需求、典型通信技术、面向角度编队控制的应用机制以及工程化难点展开论述，"
    "分析自组网、时延补偿、事件触发通信、协同定位与抗干扰设计在编队系统中的作用，"
    "并给出适用于无人机编队通信与控制一体化设计的思路。"
)

KEYWORDS = "无人机编队；角度编队；通信技术；自组网；协同控制"

SECTIONS = [
    (
        "一、引言",
        [
            "与单架无人机相比，编队系统能够通过空间分布和任务分工扩大感知范围、提升系统冗余度，并在复杂任务中获得更好的效率与鲁棒性。"
            "但编队规模一旦扩大，控制问题就不再只是飞行器自身姿态和轨迹跟踪，而是演化为多智能体之间的信息组织问题。"
            "控制器想要稳定工作，首先要知道谁与谁需要交换信息、交换何种信息、信息多快到达以及在丢包情况下如何继续保持队形。"
            "因此，通信系统不是附属模块，而是编队控制闭环中的关键组成。"
            ,
            "从当前研究趋势看，位置型编队、距离型编队和角度型编队各有优势。"
            "围绕角度约束设计的编队控制方案通常更加依赖相对方位、相邻夹角和航向关系，"
            "对全局坐标的一致性要求相对较低，适合在卫星导航受限、环境遮挡明显或任务区域动态变化的情况下使用。"
            "但这类方法对邻接信息的连续可得性更敏感，一旦通信拓扑频繁切换、时延增大或时间同步不准，"
            "就可能引起角度误差积累、队形扭曲甚至局部碰撞。"
            "因此，研究无人机编队中通信技术的具体应用，对提升角度编队控制方案的可实现性具有直接意义。"
            ,
        ],
    ),
    (
        "二、无人机编队控制对通信系统的基本需求",
        [
            "角度编队控制首先要求通信链路支持低时延和较稳定的刷新频率。"
            "在角度约束控制中，控制量往往由相邻无人机的相对方位、速度方向或参考航迹共同决定。"
            "如果状态广播周期过长，控制器使用的就是过期信息，编队会表现出明显的滞后振荡；"
            "如果链路抖动较大，则不同无人机对同一时刻队形的理解不一致，容易出现局部拉扯现象。"
            "因此，通信系统至少要保证关键控制信息具备确定性更强的传输时序。"
            ,
            "其次，角度编队对拓扑结构和可达性有明确要求。"
            "编队控制图不仅决定信息流方向，也决定控制系统是否具备足够的角度刚性。"
            "当某一条关键边失效时，虽然网络层面可能仍然连通，但控制层面未必还能唯一确定期望队形。"
            "这说明通信拓扑设计不能只追求平均吞吐量，而应结合控制图中的关键邻接关系进行优先级划分，"
            "保证核心测量链路、领航节点广播链路和安全间隔相关链路拥有更高的可靠性。"
            ,
            "再次，无人机编队通信还必须兼顾时间同步、位姿基准统一和多源信息融合。"
            "角度信息本身往往来自视觉、测向、UWB或惯导估计，如果时间戳对不齐，同一组测量就无法在控制层正确拼接。"
            "在存在多跳转发的情况下，节点还需要知道接收数据的生成时刻、置信度和坐标基准，"
            "否则会将传输误差误判为队形误差。"
            "所以，一套可用的编队通信系统不仅是无线链路，还应包括同步机制、数据封装规范和状态一致性维护机制。"
            ,
        ],
    ),
    (
        "三、无人机编队中的关键通信技术",
        [
            "在网络组织方式上，无人机编队常采用飞行自组网架构。"
            "这类网络具有节点高速移动、拓扑频繁变化和链路易受遮挡的特点，适合使用去中心化程度较高的组网方式。"
            "当编队规模较小、队形稳定时，可以使用近似星型或领航者广播型结构，以减少路由开销；"
            "当编队规模扩大或任务空间拉伸时，更适合使用多跳网状结构，引入中继节点保持边缘无人机与核心节点的联通。"
            "从控制角度看，组网方式应与编队控制层次相适配，而不是单纯套用通用无线网络方案。"
            ,
            "在链路接入与调度方面，时分接入、优先级队列和事件触发发送是提升编队通信效率的重要手段。"
            "对控制闭环影响最大的通常不是普通载荷数据，而是位姿状态、控制指令和碰撞预警信息。"
            "因此，网络调度应让关键控制包优先发送，并尽量压缩排队等待时间。"
            "对于变化较慢的状态量，可采用事件触发机制，仅在角度误差、相对距离或速度偏差超过阈值时发送更新，"
            "从而降低信道占用率，为真正关键的数据留出传输空间。"
            ,
            "在支撑技术层面，协同定位、链路质量评估和安全抗扰同样不可缺少。"
            "仅依靠单一 GNSS 难以满足近距离编队和遮挡环境下的稳定需求，实际系统常将 UWB、视觉测量、惯性导航和机间通信数据结合起来进行联合估计。"
            "与此同时，通信节点需要根据接收信号强度、时延、丢包率和剩余能量动态调整发送功率、转发策略和邻居选择。"
            "如果任务场景存在电磁干扰或恶意截获风险，还必须结合跳频、加密认证和异常节点剔除机制，"
            "避免通信失真直接破坏编队控制闭环。"
            ,
        ],
    ),
    (
        "四、通信技术在角度编队控制中的应用机制",
        [
            "对于基于角度约束的编队控制，通信系统首先承担参考信息分发的任务。"
            "无论采用领航跟随结构、虚拟结构方法还是一致性方法，编队成员都需要获得期望航向、目标运动趋势或邻居参考边的信息。"
            "当控制目标是维持若干相对夹角时，节点之间传输的不只是位置坐标，"
            "还包括测得的视线方向、局部坐标系下的方位角、角速度估计以及控制权重。"
            "通信系统越能保持这些信息的时序一致性，角度控制律的收敛过程就越平滑。"
            ,
            "其次，通信拓扑本身就是角度编队控制图的物理载体。"
            "在理论分析中，常用图论描述邻接关系和信息流向；在工程实现中，图中的每一条边都对应一条实际可用的通信或感知链路。"
            "如果某个节点失去与关键邻居的联系，那么控制器就无法准确构造角度误差项，"
            "原本满足刚性条件的编队也可能退化为形状不确定的松散集合。"
            "因此，通信层应具备面向控制图的健康监测能力，及时判断哪一条边失效会影响编队稳定，"
            "并通过备份链路、拓扑重构或角色切换来补偿。"
            ,
            "再次，时延与丢包补偿技术直接决定角度编队在动态环境中的稳定边界。"
            "当无人机执行转弯、绕障或围绕目标机动时，角度变化速度会显著增加，"
            "若仍以固定带宽和固定周期交换信息，就容易出现控制滞后。"
            "此时可以在通信层配合时间戳预测、缓存重排序和关键帧重传，在控制层引入观测器、预测补偿和鲁棒一致性算法，"
            "使无人机在短时失联或延迟增加时仍能利用最近邻模型维持近似队形，避免编队瞬间散开。"
            ,
            "最后，事件触发通信与分层通信机制非常适合角度编队任务。"
            "对于队形中心、领航机或承担中继任务的节点，可保持较高频率的状态广播；"
            "对于几何关系稳定的普通成员，则只在角度误差接近约束边界、邻居切换或安全距离逼近阈值时上报。"
            "这样做既减轻了信道压力，也能让网络资源向真正影响角度稳定性的链路倾斜。"
            "若再结合任务层、控制层和感知层的分层消息设计，系统就能在有限带宽下兼顾编队保持、碰撞规避和任务协同。"
            ,
        ],
    ),
    (
        "五、工程实现中的瓶颈与优化方向",
        [
            "当前无人机编队通信的首要难点仍然是高机动条件下的链路稳定性。"
            "机体姿态变化、遮挡、天线方向性和复杂地形都会使链路质量在短时间内大幅波动。"
            "对角度编队而言，这种波动带来的问题并不只是数据速率下降，更关键的是关键邻接边可能间歇性消失。"
            "因此，工程设计中应将通信冗余前置考虑，例如采用双链路并行、异构传感冗余和可切换中继策略，"
            "让控制器在一类信息源退化时仍有替代依据。"
            ,
            "第二个瓶颈是通信设计与控制设计仍然存在一定割裂。"
            "很多系统在仿真中默认理想通信，而在实装阶段才发现控制参数对时延和丢包极为敏感。"
            "更合理的做法是从任务开始就进行通信与控制协同设计，"
            "把刷新周期、分组长度、拓扑切换策略和控制增益一起纳入联合优化。"
            "对围绕角度约束的方案而言，还应明确哪些角度边是刚性保持所必需的，"
            "并据此定义链路优先级、触发阈值和故障后的重构规则。"
            ,
            "未来的发展方向可集中在智能化和融合化两个方面。"
            "一方面，可利用强化学习或自适应算法根据任务阶段动态调整通信频率、发射功率和路由策略，"
            "让网络资源与编队控制需求实时匹配；另一方面，可将 5G/6G 低空通信、星地协同链路、边缘计算与机载感知融合起来，"
            "构建更强的空地协同编队系统。"
            "当通信、感知和控制形成统一的数据闭环后，角度编队控制将不再只是局部几何保持问题，"
            "而会发展为面向任务效能的智能协同决策问题。"
            ,
        ],
    ),
    (
        "六、结论",
        [
            "总体来看，无人机编队中的通信技术不仅决定信息能否传到，更决定编队控制能否稳定、可靠且可扩展地实现。"
            "对于围绕角度约束设计的编队控制方案，通信系统需要同时满足低时延、关键链路高可靠、时间同步准确和拓扑可重构等要求。"
            "飞行自组网、事件触发通信、协同定位、时延补偿以及抗干扰安全机制，为角度编队提供了可落地的技术支撑。"
            "后续研究应继续推动通信与控制的一体化设计，使编队系统在复杂环境中保持几何稳定、任务连续和工程可实施性。"
            ,
        ],
    ),
]

REFERENCES = [
    "[1] Olfati-Saber R, Fax J A, Murray R M. Consensus and cooperation in networked multi-agent systems[J]. Proceedings of the IEEE, 2007.",
    "[2] Ren W, Beard R W. Distributed Consensus in Multi-vehicle Cooperative Control[M]. London: Springer, 2008.",
    "[3] Oh K K, Park M C, Ahn H S. A survey of multi-agent formation control[J]. Automatica, 2015.",
    "[4] Bekmezci I, Sahingoz O K, Temel S. Flying Ad-Hoc Networks (FANETs): A survey[J]. Ad Hoc Networks, 2013.",
    "[5] Gupta L, Jain R, Vaszkun G. Survey of important issues in UAV communication networks[J]. IEEE Communications Surveys & Tutorials, 2016.",
    "[6] Mozaffari M, Saad W, Bennis M, et al. A tutorial on UAVs for wireless networks: Applications, challenges, and open problems[J]. IEEE Communications Surveys & Tutorials, 2019.",
    "[7] Zhao S, Zelazo D. Bearing rigidity and its applications in multi-agent systems[J]. IEEE Control Systems Magazine, 2019.",
]


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_line_spacing(paragraph, multiplier: float = 1.5) -> None:
    paragraph.paragraph_format.line_spacing = multiplier
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def add_text_paragraph(document: Document, text: str, *, font_size: float = 12, bold: bool = False,
                       first_line_indent: float | None = 24, alignment=None) -> None:
    paragraph = document.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    if first_line_indent is not None:
        paragraph.paragraph_format.first_line_indent = Pt(first_line_indent)
    run = paragraph.add_run(text)
    run.bold = bold
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(font_size)
    set_paragraph_line_spacing(paragraph)


def add_heading_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    set_east_asia_font(run, "黑体")
    run.font.size = Pt(12)
    set_paragraph_line_spacing(paragraph)


def create_figure(fig_path: Path) -> None:
    fig_path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1400, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_big = ImageFont.load_default()
    font_small = ImageFont.load_default()

    title_box = [(80, 40), (1320, 120)]
    draw.rounded_rectangle(title_box, radius=20, outline="#2F5D7C", width=4, fill="#F4F8FB")
    draw.text((440, 70), "Angle-Formation Communication Topology", fill="#1F2D3D", font=font_big)

    center = (700, 470)
    leader_r = 75
    draw.ellipse(
        (center[0] - leader_r, center[1] - leader_r, center[0] + leader_r, center[1] + leader_r),
        fill="#FAD6A5",
        outline="#C76B29",
        width=4,
    )
    draw.text((650, 455), "Leader", fill="black", font=font_big)

    follower_points = [
        (700, 210),
        (940, 350),
        (940, 610),
        (700, 750),
        (460, 610),
        (460, 350),
    ]
    follower_colors = ["#D4E6F1", "#D5F5E3", "#FCF3CF", "#F5CBA7", "#E8DAEF", "#F9E79F"]

    for index, point in enumerate(follower_points, start=1):
        x, y = point
        draw.ellipse((x - 55, y - 55, x + 55, y + 55), fill=follower_colors[index - 1], outline="#355C7D", width=3)
        draw.text((x - 28, y - 7), f"U{index}", fill="black", font=font_big)
        draw.line((center[0], center[1], x, y), fill="#C0392B", width=4)

    ring_pairs = list(zip(follower_points, follower_points[1:] + follower_points[:1]))
    for start, end in ring_pairs:
        draw.line((start[0], start[1], end[0], end[1]), fill="#2874A6", width=3)

    labels = [
        ((740, 290), "bearing info"),
        ((860, 470), "neighbor state"),
        ((740, 650), "event trigger"),
        ((510, 640), "relay link"),
        ((370, 470), "time sync"),
        ((510, 290), "safety alert"),
    ]
    for pos, text in labels:
        draw.rounded_rectangle((pos[0] - 15, pos[1] - 15, pos[0] + 150, pos[1] + 22), radius=10, outline="#95A5A6", fill="#FBFCFC")
        draw.text(pos, text, fill="#34495E", font=font_small)

    legend_box = [(1010, 690), (1320, 830)]
    draw.rounded_rectangle(legend_box, radius=16, outline="#7F8C8D", width=2, fill="#FCFCFC")
    draw.line((1040, 730, 1100, 730), fill="#C0392B", width=4)
    draw.text((1120, 720), "leader reference", fill="#2C3E50", font=font_small)
    draw.line((1040, 780, 1100, 780), fill="#2874A6", width=3)
    draw.text((1120, 770), "neighbor exchange", fill="#2C3E50", font=font_small)

    image.save(fig_path)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    section.start_type = WD_SECTION.NEW_PAGE

    normal_style = document.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(12)


def build_document() -> tuple[Path, int]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    create_figure(FIG_PATH)

    document = Document()
    configure_document(document)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(TITLE)
    title_run.bold = True
    set_east_asia_font(title_run, "黑体")
    title_run.font.size = Pt(14)
    set_paragraph_line_spacing(title_paragraph)

    add_heading_paragraph(document, "摘要")
    add_text_paragraph(document, ABSTRACT)
    add_heading_paragraph(document, "关键词")
    add_text_paragraph(document, KEYWORDS, first_line_indent=0)

    for heading, paragraphs in SECTIONS:
        add_heading_paragraph(document, heading)
        for para in paragraphs:
            add_text_paragraph(document, para)
        if heading == "四、通信技术在角度编队控制中的应用机制":
            figure_paragraph = document.add_paragraph()
            figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            figure_paragraph.paragraph_format.first_line_indent = Pt(0)
            figure_run = figure_paragraph.add_run()
            figure_run.add_picture(str(FIG_PATH), width=Inches(5.8))
            set_paragraph_line_spacing(figure_paragraph)

            caption_paragraph = document.add_paragraph()
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_paragraph.paragraph_format.first_line_indent = Pt(0)
            caption_run = caption_paragraph.add_run("图1 角度编队通信与控制信息流示意图")
            set_east_asia_font(caption_run, "宋体")
            caption_run.font.size = Pt(10.5)
            set_paragraph_line_spacing(caption_paragraph)

    add_heading_paragraph(document, "参考文献")
    for ref in REFERENCES:
        add_text_paragraph(document, ref, first_line_indent=0)

    document.save(DOCX_PATH)

    count_text = ABSTRACT + KEYWORDS + "".join("".join(paras) for _, paras in SECTIONS)
    chinese_char_count = len(re.findall(r"[\u4e00-\u9fff]", count_text))
    return DOCX_PATH, chinese_char_count


if __name__ == "__main__":
    path, char_count = build_document()
    print(f"saved={path}")
    print(f"chinese_chars={char_count}")
