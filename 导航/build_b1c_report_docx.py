"""Build a DOCX report from report_b1c_maincode.md using the existing report as a template."""

from __future__ import annotations

from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from markdown import markdown

ROOT = Path(__file__).resolve().parent
MARKDOWN_PATH = ROOT / "report_b1c_maincode.md"
TEMPLATE_PATH = ROOT / "output" / "doc" / "B1C主码仿真报告.docx"
OUTPUT_PATH = ROOT / "output" / "doc" / "B1C主码仿真报告_整理版.docx"
IMAGE_WIDTH = Inches(6.1)
BODY_FONT_SIZE = Pt(12)
HEADING_FONT_SIZE = Pt(14)
CAPTION_FONT_SIZE = Pt(10.5)
BODY_FIRST_LINE_INDENT = Pt(24)
ASCII_FONT = "Times New Roman"
BODY_EAST_ASIA_FONT = "宋体"
HEADING_EAST_ASIA_FONT = "黑体"

HEADING_STYLE = {
    2: "Heading 1",
    3: "Heading 2",
    4: "Heading 3",
    5: "Heading 3",
    6: "Heading 3",
}


def clear_document_body(document: Document) -> None:
    """Remove all body content but keep section properties, styles, and footer/header settings."""
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_east_asia_font(run, font_name: str) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_run_fonts(run, east_asia_font: str, ascii_font: str = ASCII_FONT, size=None) -> None:
    run.font.name = ascii_font
    set_run_east_asia_font(run, east_asia_font)
    if size is not None:
        run.font.size = size


def apply_body_paragraph_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = BODY_FIRST_LINE_INDENT
    paragraph.paragraph_format.line_spacing = 1.5


def apply_heading_paragraph_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5


def apply_caption_paragraph_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5


def apply_code_paragraph_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    apply_caption_paragraph_format(paragraph)
    run = paragraph.add_run(text)
    set_run_fonts(run, HEADING_EAST_ASIA_FONT, size=HEADING_FONT_SIZE)
    run.bold = True


def add_heading(document: Document, level: int, text: str) -> None:
    paragraph = document.add_paragraph(style=HEADING_STYLE.get(level, "Heading 3"))
    apply_heading_paragraph_format(paragraph)
    run = paragraph.add_run(text)
    set_run_fonts(run, HEADING_EAST_ASIA_FONT, size=HEADING_FONT_SIZE)
    run.bold = True


def add_normal_paragraph(document: Document):
    paragraph = document.add_paragraph(style="Normal")
    apply_body_paragraph_format(paragraph)
    return paragraph


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    apply_caption_paragraph_format(paragraph)
    run = paragraph.add_run(text)
    set_run_fonts(run, BODY_EAST_ASIA_FONT, size=CAPTION_FONT_SIZE)


def add_code_block(document: Document, code: str) -> None:
    for line in code.rstrip("\n").splitlines() or [""]:
        paragraph = document.add_paragraph(style="Code")
        apply_code_paragraph_format(paragraph)
        run = paragraph.add_run(line)
        set_run_fonts(run, "Consolas", ascii_font="Consolas", size=Pt(10.5))


def write_inline(paragraph, node) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            set_run_fonts(run, BODY_EAST_ASIA_FONT, size=BODY_FONT_SIZE)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "br":
        paragraph.add_run().add_break()
        return

    if node.name == "code":
        run = paragraph.add_run(node.get_text())
        set_run_fonts(run, "Consolas", ascii_font="Consolas", size=Pt(10.5))
        return

    if node.name in {"strong", "b"}:
        run = paragraph.add_run(node.get_text())
        run.bold = True
        set_run_fonts(run, BODY_EAST_ASIA_FONT, size=BODY_FONT_SIZE)
        return

    if node.name in {"em", "i"}:
        run = paragraph.add_run(node.get_text())
        run.italic = True
        set_run_fonts(run, BODY_EAST_ASIA_FONT, size=BODY_FONT_SIZE)
        return

    if node.name == "a":
        text = node.get_text()
        if text:
            run = paragraph.add_run(text)
            run.underline = True
            set_run_fonts(run, BODY_EAST_ASIA_FONT, size=BODY_FONT_SIZE)
        return

    for child in node.children:
        write_inline(paragraph, child)


def extract_list_item_nodes(li: Tag) -> list[Tag | NavigableString]:
    items: list[Tag | NavigableString] = []
    for child in li.children:
        if isinstance(child, NavigableString):
            if str(child).strip():
                items.append(child)
            continue
        if child.name in {"ol", "ul"}:
            continue
        if child.name == "p":
            items.extend(list(child.children))
        else:
            items.append(child)
    return items


def add_list(document: Document, list_tag: Tag) -> None:
    ordered = list_tag.name == "ol"
    start = int(list_tag.get("start", 1))
    index = start
    for li in list_tag.find_all("li", recursive=False):
        paragraph = add_normal_paragraph(document)
        prefix = f"{index}. " if ordered else "• "
        run = paragraph.add_run(prefix)
        set_run_east_asia_font(run, "宋体")
        for child in extract_list_item_nodes(li):
            write_inline(paragraph, child)
        nested_lists = [child for child in li.children if isinstance(child, Tag) and child.name in {"ol", "ul"}]
        for nested in nested_lists:
            add_list(document, nested)
        index += 1


def resolve_image(src: str) -> Path:
    path = (MARKDOWN_PATH.parent / src).resolve()
    if not path.exists():
        raise FileNotFoundError(f"image not found: {src}")
    return path


def add_figure(document: Document, image_path: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    apply_caption_paragraph_format(paragraph)
    paragraph.add_run().add_picture(str(image_path), width=IMAGE_WIDTH)
    add_caption(document, caption)


def add_table(document: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr")
    if not rows:
        return

    first_row_cells = rows[0].find_all(["th", "td"])
    table = document.add_table(rows=len(rows), cols=len(first_row_cells), style="Table Grid")

    for row_index, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for col_index, cell in enumerate(cells):
            target = table.cell(row_index, col_index)
            target.text = cell.get_text(" ", strip=True)
            for paragraph in target.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.5
                for run in paragraph.runs:
                    set_run_fonts(run, BODY_EAST_ASIA_FONT, size=BODY_FONT_SIZE)
                    if row_index == 0:
                        run.bold = True


def build_document() -> Document:
    document = Document(str(TEMPLATE_PATH))
    clear_document_body(document)
    return document


def render_markdown(document: Document, text: str) -> None:
    html = markdown(text, extensions=["fenced_code", "tables"])
    soup = BeautifulSoup(html, "html.parser")
    figure_index = 1

    for node in soup.children:
        if isinstance(node, NavigableString) and not str(node).strip():
            continue
        if not isinstance(node, Tag):
            continue

        if node.name == "h1":
            add_title(document, node.get_text(strip=True))
            continue

        if node.name in {"h2", "h3", "h4", "h5", "h6"}:
            add_heading(document, int(node.name[1]), node.get_text(strip=True))
            continue

        if node.name == "p":
            image = node.find("img", recursive=False)
            tags = [child for child in node.children if isinstance(child, Tag)]
            if image and len(tags) == 1 and tags[0].name == "img" and not node.get_text(strip=True):
                caption = image.get("alt", "图片")
                add_figure(document, resolve_image(image["src"]), f"图 {figure_index} {caption}")
                figure_index += 1
                continue

            paragraph = add_normal_paragraph(document)
            for child in node.children:
                write_inline(paragraph, child)
            continue

        if node.name in {"ol", "ul"}:
            add_list(document, node)
            continue

        if node.name == "pre":
            code = node.get_text()
            add_code_block(document, code)
            continue

        if node.name == "table":
            add_table(document, node)
            continue


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    markdown_text = MARKDOWN_PATH.read_text(encoding="utf-8")
    document = build_document()
    render_markdown(document, markdown_text)
    document.save(str(OUTPUT_PATH))
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
