import json
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
MD_PATH = BASE_DIR / "assignment2_report.md"
DOCX_PATH = BASE_DIR / "assignment2_report.docx"
OUTPUT_DIR = BASE_DIR / "output"
SUMMARY_PATH = OUTPUT_DIR / "summary.json"

FONT_SONGTI = "\u5b8b\u4f53"
FONT_HEITI = "\u9ed1\u4f53"
FONT_CODE = "Courier New"
CODE_SIZE_PT = 10.5

INLINE_IMAGES = {
    "3 \u7a0b\u5e8f\u5b9e\u73b0": [
        {"type": "image", "caption": "\u56fe1 X1\u3001\u53c2\u8003 X2 \u4ee5\u53ca C1 \u524d 128 \u4e2a\u7801\u7247", "filename": "first_128_chips.png"},
    ],
    "4 \u5b9e\u9a8c\u7ed3\u679c\u5206\u6790": [
        {"type": "image", "caption": "\u56fe2 \u5b8c\u6574 X1/X2 \u4f18\u9009\u5bf9\u5468\u671f\u4e92\u76f8\u5173", "filename": "preferred_pair_crosscorr.png"},
        {"type": "image", "caption": "\u56fe3 \u4ee3\u8868\u6027\u77ed P \u7801\u5e8f\u5217 C1 \u7684\u5468\u671f\u81ea\u76f8\u5173", "filename": "one_second_autocorr.png"},
        {"type": "table", "table": "crosscorr_summary", "caption": "\u88682 \u77ed P \u7801\u5e8f\u5217\u4e92\u76f8\u5173\u7edf\u8ba1\u7ed3\u679c"},
        {"type": "image", "caption": "\u56fe4 \u6700\u5dee\u7801\u5bf9 C1-C2 \u7684\u5468\u671f\u4e92\u76f8\u5173\u66f2\u7ebf", "filename": "one_second_crosscorr_example.png"},
    ],
}


def sanitize_text(text: str) -> str:
    return text.replace("`", "")


def set_run_font(run, size_pt: float, east_asia: str = FONT_SONGTI, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_code_font(run, size_pt: float = CODE_SIZE_PT) -> None:
    run.font.name = FONT_CODE
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CODE)
    run.font.size = Pt(size_pt)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)


def style_paragraph(para, *, heading: bool = False, title: bool = False) -> None:
    fmt = para.paragraph_format
    fmt.line_spacing = 1.5
    fmt.first_line_indent = Pt(0) if heading or title else Pt(24)
    if title:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif heading:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for run in para.runs:
        if title or heading:
            set_run_font(run, 14, east_asia=FONT_HEITI, bold=True)
        else:
            set_run_font(run, 12)


def add_paragraph(doc: Document, text: str, *, heading: bool = False, title: bool = False) -> None:
    para = doc.add_paragraph()
    para.add_run(sanitize_text(text))
    style_paragraph(para, heading=heading, title=title)


def add_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.add_run(sanitize_text(text))
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.first_line_indent = Pt(0)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        set_run_font(run, 12)


def add_code_block(doc: Document, lines: list[str]) -> None:
    para = doc.add_paragraph()
    fmt = para.paragraph_format
    fmt.line_spacing = 1.0
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(18)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run("\n".join(lines))
    set_code_font(run)


def add_register2_states_table(doc: Document) -> None:
    summary = json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))
    family = summary["initial_register_states"]["registers2_family"]
    labels = list(family.keys())
    table = doc.add_table(rows=len(labels) + 1, cols=3)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "序号"
    header_cells[1].text = "码序列"
    header_cells[2].text = "registers2 初始状态"

    for i, label in enumerate(labels, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = str(i)
        row_cells[1].text = label
        row_cells[2].text = "[" + " ".join(str(v) for v in family[label]) + "]"

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.line_spacing = 1.2
                para.paragraph_format.first_line_indent = Pt(0)
                for run in para.runs:
                    set_run_font(run, 10.5, bold=False)


def add_crosscorr_summary_table(doc: Document) -> None:
    summary = json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))
    cc = summary["one_second_correlation"]["crosscorrelation"]
    rows = [
        ("零延迟互相关", "平均绝对值", "-", f"{cc['mean_pair_abs_zero_lag']:.6f}"),
        (
            "零延迟互相关",
            "最差码对",
            "-".join(cc["worst_zero_lag_pair"]["pair"]),
            f"{cc['worst_zero_lag_pair']['zero_lag']:.6f}",
        ),
        ("全延迟最大互相关", "平均 max|R|", "-", f"{cc['mean_pair_max_abs']:.6f}"),
        (
            "全延迟最大互相关",
            "最差码对",
            "-".join(cc["worst_search_pair"]["pair"]),
            f"{cc['worst_search_pair']['max_abs']:.6f}",
        ),
    ]

    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = "Table Grid"
    headers = ["指标类别", "统计项", "码对", "数值"]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text

    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            table.rows[r].cells[c].text = value

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.line_spacing = 1.2
                para.paragraph_format.first_line_indent = Pt(0)
                for run in para.runs:
                    set_run_font(run, 10.5, bold=False)


def flush_paragraph(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(part.strip() for part in buffer if part.strip()).strip()
    if text:
        add_paragraph(doc, text)
    buffer.clear()


def insert_inline_images(doc: Document, heading_text: str) -> None:
    clean_heading = sanitize_text(heading_text)
    for item in INLINE_IMAGES.get(clean_heading, []):
        if item["type"] == "image":
            path = OUTPUT_DIR / item["filename"]
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.2))
                add_caption(doc, item["caption"])
        elif item["type"] == "table" and SUMMARY_PATH.exists():
            if item["table"] == "crosscorr_summary":
                add_crosscorr_summary_table(doc)
                add_caption(doc, item["caption"])
    if clean_heading == "2 理论基础" and SUMMARY_PATH.exists():
        add_register2_states_table(doc)
        add_caption(doc, "表1 10组 registers2 初始状态")


def build_docx(output_path: Path | None = None) -> None:
    if output_path is None:
        output_path = DOCX_PATH

    text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(85.05)
    section.right_margin = Pt(70.85)

    paragraph_buffer: list[str] = []
    code_block_buffer: list[str] = []
    in_code_block = False

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                add_code_block(doc, code_block_buffer)
                code_block_buffer.clear()
                in_code_block = False
            else:
                flush_paragraph(doc, paragraph_buffer)
                in_code_block = True
            continue

        if in_code_block:
            code_block_buffer.append(line)
            continue

        if not stripped:
            flush_paragraph(doc, paragraph_buffer)
            continue

        if stripped.startswith("# "):
            flush_paragraph(doc, paragraph_buffer)
            add_paragraph(doc, stripped[2:].strip(), title=True)
            continue

        if stripped.startswith("## "):
            flush_paragraph(doc, paragraph_buffer)
            heading_text = stripped[3:].strip()
            add_paragraph(doc, heading_text, heading=True)
            insert_inline_images(doc, heading_text)
            continue

        if stripped.startswith("### "):
            flush_paragraph(doc, paragraph_buffer)
            heading_text = stripped[4:].strip()
            add_paragraph(doc, heading_text, heading=True)
            insert_inline_images(doc, heading_text)
            continue

        if stripped.startswith("- "):
            flush_paragraph(doc, paragraph_buffer)
            add_paragraph(doc, stripped[2:].strip())
            continue

        numbered = False
        for idx in range(1, 20):
            prefix = f"{idx}. "
            if stripped.startswith(prefix):
                flush_paragraph(doc, paragraph_buffer)
                add_paragraph(doc, stripped)
                numbered = True
                break
        if numbered:
            continue

        if stripped.startswith("`") and stripped.endswith("`"):
            flush_paragraph(doc, paragraph_buffer)
            add_paragraph(doc, stripped.strip("`"))
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph(doc, paragraph_buffer)
    if code_block_buffer:
        add_code_block(doc, code_block_buffer)
    doc.save(output_path)


if __name__ == "__main__":
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else DOCX_PATH
    build_docx(target)
